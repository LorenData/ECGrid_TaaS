# mapping-script
import io
import bots.transform as transform

from docx import Document
from docx.shared import Inches


def main(inn, out):

    out.ta_info['frompartner'] = ""
    out.ta_info['topartner'] = ""
    out.ta_info['testindicator'] = "P"
    inn.ta_info['testindicator'] = "P"
    inn.ta_info['botskey'] = "HDGS1-128"
    out.ta_info['botskey'] = "HDGS1-128"

    wordData = io.BytesIO()

    document = Document()

    document.add_heading('Document Title', 0)

    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    reference = inn.get({'BOTSID': 'ST'}, {'BOTSID': 'BEG', 'BEG03': None})
    orderdate = inn.get({'BOTSID': 'ST'}, {'BOTSID': 'BEG', 'BEG05': None})
    purpose_code = inn.get({'BOTSID': 'ST'}, {'BOTSID': 'BEG', 'BEG01': None})
    SCAC = inn.get({'BOTSID': 'ST'}, {'BOTSID': 'TD5', 'TD502': '2', 'TD503': None})
    CarrierName = inn.get({'BOTSID': 'ST'}, {'BOTSID': 'TD5', 'TD505': None})

    z = "Reference: %s and Order Data: %s" % (reference,orderdate)

    document.add_paragraph(z)

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')

    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )

    #document.add_picture('', width=Inches(1.25))

    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    document.add_page_break()

    document.save(wordData)

    # Write to Output.
    out.root = wordData.getvalue()

    wordData.close()
